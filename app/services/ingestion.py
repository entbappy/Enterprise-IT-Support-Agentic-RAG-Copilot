from pathlib import Path
from typing import Iterable
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument


SUPPORTED = {".pdf", ".txt", ".md", ".docx"}



def load_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return PyPDFLoader(str(path)).load()
    if suffix in {".txt", ".md"}:
        return TextLoader(str(path), encoding="utf-8").load()
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"source": str(path)})]
    raise ValueError(f"Unsupported file type: {suffix}")



def chunk_documents(docs: Iterable[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=120, add_start_index=True)
    return splitter.split_documents(list(docs))
